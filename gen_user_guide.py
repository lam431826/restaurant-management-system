"""
Generate WRMS User Guide — chỉ ghi các tính năng đã kết nối BE thật sự.
Run: python gen_user_guide.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x02, 0x5c, 0xca)
DARK      = RGBColor(0x20, 0x23, 0x25)
MUTED     = RGBColor(0x6b, 0x72, 0x80)
SUCCESS   = RGBColor(0x0d, 0x9e, 0x6e)
DANGER    = RGBColor(0xe5, 0x39, 0x35)
ORANGE    = RGBColor(0xe6, 0x7e, 0x00)
PURPLE    = RGBColor(0x7c, 0x3a, 0xed)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLU = RGBColor(0xe8, 0xf0, 0xfe)
LIGHT_GRY = RGBColor(0xf5, 0xf6, 0xf7)

# ── XML helpers ───────────────────────────────────────────────────────────────

def shd(cell, hex6: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex6.upper().lstrip('#'))
    tcPr.append(s)


def no_space_para(cell):
    """Remove default paragraph spacing inside a table cell."""
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    return p


def run(para, text, bold=False, italic=False, color=None, size=Pt(10.5)):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = size
    if color: r.font.color.rgb = color
    return r


# ── Paragraph builders ────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = PRIMARY
    # bottom border
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '025CCA')
    pBdr.append(b); pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRIMARY


def body(doc, text, color=None, bold=False, sa=Pt(4)):
    p = doc.add_paragraph(); p.paragraph_format.space_after = sa
    r = p.add_run(text); r.font.size = Pt(10.5); r.bold = bold
    if color: r.font.color.rgb = color


def bul(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 + level * 14)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)


def note(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(6)
    c = color or PRIMARY
    r = p.add_run('  ' + text); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = c


def api_box(doc, rows):
    """Single-column table for API call documentation."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'
    t.autofit = False
    t.columns[0].width = Cm(3.5)
    t.columns[1].width = Cm(11.5)
    for i, (label, val) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        shd(c0, 'EBF0FB' if i % 2 == 0 else 'F5F6F7')
        shd(c1, 'EBF0FB' if i % 2 == 0 else 'F5F6F7')
        p0 = no_space_para(c0); r0 = p0.add_run(label)
        r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = MUTED
        p1 = no_space_para(c1); r1 = p1.add_run(val)
        r1.font.size = Pt(9.5); r1.font.name = 'Consolas'
    doc.add_paragraph()


def step_tbl(doc, steps):
    """Numbered step table."""
    t = doc.add_table(rows=len(steps), cols=2)
    t.style = 'Table Grid'; t.autofit = False
    t.columns[0].width = Cm(1.1); t.columns[1].width = Cm(13.9)
    for i, (num, desc) in enumerate(steps):
        c0, c1 = t.rows[i].cells
        shd(c0, 'E8F0FE'); shd(c1, 'FFFFFF' if i % 2 == 0 else 'F9FAFF')
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = no_space_para(c0); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(num); r0.bold = True; r0.font.color.rgb = PRIMARY; r0.font.size = Pt(11)
        p1 = no_space_para(c1); r1 = p1.add_run(desc); r1.font.size = Pt(10.5)
    doc.add_paragraph()


def role_line(doc, roles):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    run(p, 'Vai tro: ', bold=True, color=MUTED, size=Pt(9.5))
    label_colors = {'ADMIN': PURPLE, 'MANAGER': PRIMARY, 'CASHIER': SUCCESS, 'WAITER': ORANGE}
    for r in roles:
        c = label_colors.get(r, MUTED)
        rn = p.add_run(f' [{r}] '); rn.bold = True; rn.font.size = Pt(9.5); rn.font.color.rgb = c
        p.add_run('  ')


def page_break(doc):
    doc.add_page_break()


SHOTS_DIR = 'd:/SU26/SWP/restaurant-management-system/docs/screenshots'


def add_screenshot(doc, filename, caption=None, width=Cm(14)):
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(8)
        rc = cp.add_run(caption)
        rc.font.size = Pt(9)
        rc.italic = True
        rc.font.color.rgb = MUTED


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
sec.top_margin  = Cm(2.5); sec.bottom_margin = Cm(2.0)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ── Cover ─────────────────────────────────────────────────────────────────────
cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(70)
r = cover.add_run('WASABI RESTAURANT MANAGEMENT SYSTEM')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = PRIMARY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(12)
rs = sub.add_run('Tai lieu Huong dan Su dung he thong quan ly (Staff Portal)')
rs.bold = True; rs.font.size = Pt(15); rs.font.color.rgb = DARK

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.paragraph_format.space_before = Pt(20)
rv = ver.add_run('Phien ban 1.1  -  Thang 6 / 2026  |  Chi tai lieu tinh nang da ket noi Backend')
rv.font.size = Pt(10); rv.font.color.rgb = MUTED; rv.italic = True

page_break(doc)

# ── Scope note ────────────────────────────────────────────────────────────────
h1(doc, 'PHAM VI TAI LIEU')
body(doc,
     'Tai lieu nay chi mo ta cac tinh nang THUC SU KET NOI toi Backend API (khong bao gom man hinh dung du lieu '
     'mau/mockdata). Moi tinh nang duoi day da duoc xac nhan co API call that trong source code.',
     color=DARK)

body(doc, 'Cac tinh nang da co UI nhung CHUA ket noi BE (dung mockData):')
bul(doc, 'Dashboard KPI / Bieu do doanh thu / Bieu do khach hang -- dung kpiData, chartData tu mockData.ts')
bul(doc, 'Quan ly Thuc don (/manager/products) -- dung products[] tu mockData.ts')
bul(doc, 'Quan ly Khu vuc & Ban (/manager/rooms) -- dung rooms[] tu mockData.ts')
bul(doc, 'Quan ly Nhan vien (/manager/employees) -- dung employees[] tu mockData.ts')
bul(doc, 'Hoa don & Giao dich (/manager/invoices) -- dung invoices[] tu mockData.ts')
bul(doc, 'Cashier: tao order, danh sach mon -- MENU_ITEMS la mang hardcode; order CHUA gui toi API')
body(doc, '', sa=Pt(2))
body(doc, 'Cac tinh nang DA KET NOI BE (toan bo noi dung trong tai lieu nay):')
bul(doc, 'Xac thuc (Login / OTP / Forgot / Reset / Change password / Logout)')
bul(doc, 'Quan ly Dat ban (CRUD + Confirm / Cancel / Check-in / No-show / Assign table)')
bul(doc, 'Nhat ky thao tac (Audit Log)')
bul(doc, 'Hoat dong gan day tren Dashboard (widget RecentActivities)')
bul(doc, 'Cashier: lay danh sach ban tu API -- GET /api/tables')
bul(doc, 'Thong bao email (poll ket qua gui email sau khi thao tac dat ban)')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. XAC THUC
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '1. Xac thuc (Authentication)')
role_line(doc, ['ADMIN', 'MANAGER', 'CASHIER', 'WAITER'])
body(doc,
     'Tat ca API xac thuc dung qua src/api/auth.ts goi toi apiClient (axios, baseURL=/api). '
     'Token duoc luu vao localStorage: access_token, refresh_token, user (JSON). '
     'apiClient tu dong dinh kem Authorization: Bearer <token> moi request.')

h2(doc, '1.1  Dang nhap')
body(doc, 'File: src/components/auth/LoginPage.tsx', color=MUTED)
body(doc, 'Luong 3 buoc (state machine: login -> send-otp -> enter-otp):')

h3(doc, 'Buoc 1 -- Nhap tai khoan mat khau (tai khoan ACTIVE)')
step_tbl(doc, [
    ('1', 'Nguoi dung nhap username + password, nhan "Dang nhap".'),
    ('2', 'Goi: POST /api/auth/login   body: { username, password }'),
    ('3', 'Response tra ve { accessToken, refreshToken, user: {id, username, fullName, role}, requiresVerification: false }'),
    ('4', 'saveSession() luu token vao localStorage.'),
    ('5', 'navigate(defaultRoute(role)) -- ADMIN -> /admin, MANAGER -> /manager/dashboard, CASHIER -> /cashier, WAITER -> /waiter'),
])
api_box(doc, [
    ('Method', 'POST'),
    ('Endpoint', '/api/auth/login'),
    ('Body', '{ "username": "...", "password": "..." }'),
    ('Response 200', '{ accessToken, refreshToken, user: {id,username,fullName,role}, requiresVerification, verifyToken }'),
    ('Error 401', 'Sai username/password -- hien "Ten dang nhap hoac mat khau khong dung."'),
    ('Error 423', 'Tai khoan bi khoa -- hien "Tai khoan dang bi khoa."'),
])

h3(doc, 'Buoc 1b -- Tai khoan UN_ACTIVE (dang nhap lan dau)')
body(doc,
     'Neu response co requiresVerification: true, hien thi buoc xac thuc OTP. '
     'verifyToken tra ve duoc luu vao state.')

step_tbl(doc, [
    ('2', 'Nguoi dung nhan "Gui ma OTP".'),
    ('3', 'Goi: POST /api/auth/verify/info   header: X-Verify-Token: <verifyToken>'),
    ('4', 'Server gui OTP 6 chu so den email cua tai khoan, tra ve { maskedEmail, expiresAt }.'),
    ('5', 'Nguoi dung nhap OTP vao o xac nhan, nhan "Xac nhan".'),
    ('6', 'Goi: POST /api/auth/verify/otp   header: X-Verify-Token   body: { otp }'),
    ('7', 'Thanh cong: tai khoan chuyen ACTIVE, server tra JWT day du. saveSession() -> navigate.'),
])
api_box(doc, [
    ('verify/info Method', 'POST'),
    ('verify/info Endpoint', '/api/auth/verify/info'),
    ('verify/info Header', 'X-Verify-Token: <verifyToken>'),
    ('verify/info Response', '{ maskedEmail: "ha***@gmail.com", expiresAt: "..." }'),
    ('verify/otp Method', 'POST'),
    ('verify/otp Endpoint', '/api/auth/verify/otp'),
    ('verify/otp Header', 'X-Verify-Token: <verifyToken>'),
    ('verify/otp Body', '{ "otp": "123456" }'),
    ('verify/otp Response 200', 'JWT day du (accessToken, refreshToken, user)'),
    ('Error 401 otp sai', 'Hien thong bao OTP sai'),
    ('Error 429 max attempt', 'Qua 5 lan sai OTP'),
])

note(doc, 'Gui lai OTP: POST /api/auth/resend-otp   body: { verifyToken }   -- toi da 3 lan trong 10 phut.', color=ORANGE)
add_screenshot(doc, '01_login.png', 'Man hinh dang nhap WRMS')

h2(doc, '1.2  Quen mat khau')
body(doc, 'File: ForgotPasswordPage.tsx -> NewPasswordPage.tsx', color=MUTED)
step_tbl(doc, [
    ('1', 'Nguoi dung nhap username, nhan "Tiep tuc".'),
    ('2', 'Goi: POST /api/auth/forgot-password   body: { username }'),
    ('3', 'Server gui OTP qua email, tra ve { resetToken, maskedEmail, expiresAt }.'),
    ('4', 'navigate("/new-password", state: { resetToken, maskedEmail })'),
    ('5', 'Nguoi dung nhap OTP + mat khau moi + xac nhan mat khau, nhan "Dat lai mat khau".'),
    ('6', 'Goi: POST /api/auth/reset-password   header: X-Reset-Token: <resetToken>   body: { otp, newPassword }'),
    ('7', 'Thanh cong (HTTP 204) -> navigate("/login", { message: "Dat lai mat khau thanh cong" })'),
])
api_box(doc, [
    ('forgot-password Body', '{ "username": "username_waiter" }'),
    ('forgot-password Response', '{ resetToken, maskedEmail, expiresAt }'),
    ('reset-password Header', 'X-Reset-Token: <resetToken>'),
    ('reset-password Body', '{ "otp": "123456", "newPassword": "NewPass@123" }'),
    ('reset-password Response', 'HTTP 204 No Content'),
    ('Error 401', 'OTP sai hoac het han'),
    ('Error 429', 'Nhap sai OTP qua 5 lan'),
])

add_screenshot(doc, '02_forgot_password.png', 'Man hinh quen mat khau')

h2(doc, '1.3  Doi mat khau (khi dang nhap)')
body(doc, 'File: ChangePasswordModal.tsx -- mo tu avatar menu (tren Header WAITER/CASHIER hoac Layout MANAGER).', color=MUTED)
step_tbl(doc, [
    ('1', 'Nhan vao ten dang nhap / avatar goc tren phai -> chon "Doi mat khau".'),
    ('2', 'Nhap mat khau hien tai va mat khau moi (>= 8 ky tu, khong trung mat khau cu).'),
    ('3', 'Goi: POST /api/auth/change-password   body: { currentPassword, newPassword }'),
    ('4', 'Thanh cong: hien checkmark xanh, modal tu dong dong.'),
])
api_box(doc, [
    ('Endpoint', 'POST /api/auth/change-password'),
    ('Header', 'Authorization: Bearer <access_token> (tu dong gan qua apiClient)'),
    ('Body', '{ "currentPassword": "...", "newPassword": "..." }'),
    ('Response', 'HTTP 204'),
    ('Error 401', 'Mat khau hien tai sai'),
])

h2(doc, '1.4  Dang xuat')
body(doc, 'File: Reservation.tsx (handleLogout) / CashierOrders.tsx -- nut Logout trong avatar menu.', color=MUTED)
step_tbl(doc, [
    ('1', 'Nguoi dung nhan "Dang xuat".'),
    ('2', 'Goi: POST /api/auth/logout   (thu hoi refresh token phia server)'),
    ('3', 'signOut() xoa localStorage, navigate("/login").'),
])
api_box(doc, [
    ('Endpoint', 'POST /api/auth/logout'),
    ('Header', 'Authorization: Bearer <access_token>'),
    ('Response', 'HTTP 204'),
    ('Luu y', 'Neu request loi (token da het han), van goi signOut() va redirect login.'),
])

h2(doc, '1.5  Tu dong lam moi token (401 retry)')
body(doc,
     'apiClient.ts co response interceptor: khi nhan HTTP 401, tu dong goi '
     'POST /api/auth/refresh   body: { refreshToken }   roi thu lai request goc 1 lan. '
     'Neu refresh that bai -> signOut() + redirect /login.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. QUAN LY DAT BAN
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '2. Quan ly Dat ban')
role_line(doc, ['WAITER', 'MANAGER'])
body(doc,
     'Route: /waiter (WAITER / MANAGER). '
     'Component chinh: src/components/reservation/Reservation.tsx. '
     'API module: src/api/reservations.ts, src/api/tables.ts, src/api/notifications.ts')

h2(doc, '2.1  Khoi dong -- Load du lieu')
body(doc, 'Khi component mount, chay 2 call song song (useEffect):')
api_box(doc, [
    ('Load dat ban', 'GET /api/reservations?page=0&size=200'),
    ('Response path', 'res.data.data.data[] -- ApiResponse<PageResponse<ReservationDto>>'),
    ('Load ban', 'GET /api/tables'),
    ('Response path', 'res.data.data[] -- ApiResponse<TableDto[]>'),
    ('Muc dich ban', 'Dung cho dropdown Xep ban (CalendarView AssignTableDropdown)'),
])
body(doc, 'ReservationDto cac truong: id, tableId, tableName, tableArea, guestName, phone, guestEmail, partySize, datetime (ISO), note, status, createdAt')
body(doc, 'TableDto cac truong: id, name, capacity, area, status (AVAILABLE/OCCUPIED/RESERVED/BILLING/CLEANING), qrToken')

h2(doc, '2.2  Giao dien hien thi')

h3(doc, 'Calendar View (tab mac dinh)')
bul(doc, 'Timeline ngang 14:00 -- 23:00, truc doc la danh sach dat ban trong ngay dang chon.')
bul(doc, 'Mini Calendar trai: chon ngay, chuyen thang bang nut <- ->. Ngay hom nay co nen xanh nhat.')
bul(doc, 'Checkbox loc trang thai: PENDING / CONFIRMED / CHECKED_IN / NO_SHOW / CANCELLED.')
bul(doc, 'Nhan vao thanh mau cua 1 dat ban: mo panel chi tiet ben duoi timeline.')
bul(doc, 'Panel chi tiet hien thi: ten khach, SDT, so khach, ban, gio den, ghi chu, cac nut hanh dong.')

add_screenshot(doc, '03_reservation_calendar.png', 'Giao dien Calendar View -- dat ban theo lich')

h3(doc, 'List View')
bul(doc, 'Bang hien thi toan bo dat ban trong ngay chon: Ma / Gio den / Khach / SDT / So khach / Ban / Trang thai / Ghi chu.')
bul(doc, 'Nhan vao hang de mo panel chi tiet ben duoi bang.')
add_screenshot(doc, '04_reservation_list.png', 'Giao dien List View -- dat ban theo danh sach')

h2(doc, '2.3  Tao dat ban (F1 / nut "+ Dat ban")')
body(doc, 'File: ReservationModal.tsx', color=MUTED)
body(doc,
     'Modal mo khi nhan nut "Dat ban (F1)". Du lieu ban trong modal duoc lay tu '
     'GET /api/tables khi modal mount (chi hien AVAILABLE, nhom theo area).')
step_tbl(doc, [
    ('1', 'Nhap: Ten khach (*), SDT (* phai bat dau bang 0, 10-11 so), Email khach (tuy chon), So khach (*), Ngay & Gio den (*), Chon ban (tuy chon -- chi hien ban AVAILABLE), Ghi chu.'),
    ('2', 'Client validate: ten, SDT, so khach >= 1, email dung dinh dang neu co.'),
    ('3', 'Goi: POST /api/reservations   body: { guestName, phone, partySize, datetime ("YYYY-MM-DDTHH:mm:ss"), tableId|null, note|null, guestEmail|null }'),
    ('4', 'Thanh cong: modal dong, goi lai load() de refresh danh sach.'),
])
api_box(doc, [
    ('Endpoint', 'POST /api/reservations'),
    ('Body vi du', '{ "guestName":"Nguyen Van A", "phone":"0912345678", "partySize":4, "datetime":"2026-06-25T19:00:00", "tableId":null, "note":null, "guestEmail":"a@gmail.com" }'),
    ('Response 200', 'ApiResponse<ReservationDto>  -- res.data.data'),
    ('Luu y quan trong', 'Dat ban tao boi staff co trang thai mac dinh la CONFIRMED (khac online dat ban la PENDING)'),
    ('Error 400', 'fieldErrors: thong bao loi cu the tren tung truong'),
    ('Error 409', 'Ban da duoc xep cho dat ban khac'),
    ('Error 422', 'So khach vuot suc chua cua ban'),
])
add_screenshot(doc, '06_create_reservation_modal.png', 'Modal tao dat ban moi')

h2(doc, '2.4  Xac nhan dat ban (PENDING -> CONFIRMED)')
body(doc, 'Ap dung khi dat ban online (trang thai PENDING). Dat ban do staff tao tu dong la CONFIRMED.', color=ORANGE, bold=False)
step_tbl(doc, [
    ('1', 'Chon dat ban PENDING tu Calendar / List.'),
    ('2', 'Nhan nut "Xac nhan" trong panel chi tiet.'),
    ('3', 'Goi: PUT /api/reservations/{id}/confirm'),
    ('4', 'Thanh cong: goi load(), hien toast "Da xac nhan".'),
    ('5', 'Neu khach co email: hien toast "Dang gui email...", bat dau pollEmailResult(id).'),
])
api_box(doc, [
    ('Endpoint', 'PUT /api/reservations/{id}/confirm'),
    ('Body', '(khong can body)'),
    ('Response', 'ApiResponse<ReservationDto>  -- trang thai moi: CONFIRMED'),
    ('Side effect', 'BE gui email thong bao xac nhan cho khach (async)'),
])

h2(doc, '2.5  Xep ban (Assign Table)')
body(doc, 'Co the thuc hien khi dat ban o trang thai PENDING hoac CONFIRMED.', color=MUTED)
step_tbl(doc, [
    ('1', 'Chon dat ban can xep ban.'),
    ('2', 'Mo panel chi tiet -> nhan dropdown "Xep ban".'),
    ('3', 'Dropdown loc: chi hien ban co status AVAILABLE va capacity >= partySize (so khach).'),
    ('4', 'Chon ban mong muon.'),
    ('5', 'Goi: PUT /api/reservations/{id}   body: { tableId: "uuid-cua-ban" }'),
    ('6', 'Thanh cong: load lai, toast "Da xep ban thanh cong".'),
])
api_box(doc, [
    ('Endpoint', 'PUT /api/reservations/{id}'),
    ('Body', '{ "tableId": "7f3c1a2b-..." }'),
    ('Response', 'ApiResponse<ReservationDto>'),
    ('Error 422', 'So khach ({partySize}) vuot suc chua cua ban ({capacity}) -- TABLE_CAPACITY_EXCEEDED'),
    ('Error 404', 'Khong tim thay dat ban hoac ban'),
    ('Loc dropdown FE', 'tables.filter(t => t.status === "AVAILABLE" && t.capacity >= partySize)'),
])

h2(doc, '2.6  Check-in (CONFIRMED -> CHECKED_IN)')
step_tbl(doc, [
    ('1', 'Chon dat ban CONFIRMED.'),
    ('2', 'Nhan nut "Check-in" trong panel chi tiet.'),
    ('3', 'Goi: PUT /api/reservations/{id}/check-in'),
    ('4', 'Thanh cong: load lai, toast "Check-in thanh cong".'),
])
api_box(doc, [
    ('Endpoint', 'PUT /api/reservations/{id}/check-in'),
    ('Response', 'ApiResponse<ReservationDto>  -- trang thai moi: CHECKED_IN'),
])

h2(doc, '2.7  Danh dau Khong den (CONFIRMED -> NO_SHOW)')
step_tbl(doc, [
    ('1', 'Chon dat ban CONFIRMED.'),
    ('2', 'Nhan nut "Khong den (No-show)" trong panel chi tiet.'),
    ('3', 'Goi: PUT /api/reservations/{id}/no-show'),
    ('4', 'Thanh cong: load lai, toast "Da danh dau khong den".'),
])
api_box(doc, [
    ('Endpoint', 'PUT /api/reservations/{id}/no-show'),
    ('Response', 'HTTP 204'),
])

h2(doc, '2.8  Huy dat ban (-> CANCELLED)')
step_tbl(doc, [
    ('1', 'Chon dat ban PENDING hoac CONFIRMED.'),
    ('2', 'Nhan nut "Huy" trong panel chi tiet.'),
    ('3', 'Goi: DELETE /api/reservations/{id}'),
    ('4', 'Thanh cong: load lai, toast "Da huy dat ban".'),
    ('5', 'Neu khach co email: hien toast "Dang gui email huy...", bat dau pollEmailResult(id).'),
])
api_box(doc, [
    ('Endpoint', 'DELETE /api/reservations/{id}'),
    ('Response', 'HTTP 204'),
    ('Side effect', 'BE gui email thong bao huy cho khach (async)'),
])

h2(doc, '2.9  Xuat CSV (client-side)')
body(doc,
     'Nut "Xuat file" nhan vao goi ham exportCsv(items) -- items la mang ReservationDto da load tu API. '
     'Khong goi them API. Tao Blob CSV, tao the <a> download, click auto. '
     'Ten file: dat-ban-YYYY-MM-DD.csv. Cac cot: Ma dat ban, Gio den, Khach hang, Dien thoai, So khach, Trang thai, Ghi chu.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. THONG BAO EMAIL (NOTIFICATION POLLING)
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '3. Thong bao ket qua gui Email')
role_line(doc, ['WAITER', 'MANAGER'])
body(doc,
     'File: src/api/notifications.ts. Sau khi confirm / huy dat ban (neu khach co email), '
     'FE poll ket qua gui email async. SMTP gui email chay trong thread rieng @Async phia BE, '
     'nen FE phai cho roi hoi ket qua.')

h2(doc, '3.1  Luong Poll')
step_tbl(doc, [
    ('1', 'Sau confirm/huy: hien toast mau xanh "Dang gui email cho khach..." (10 giay).'),
    ('2', 'Sau 3 giay: Goi GET /api/notifications/log?referenceId={reservationId}&size=1'),
    ('3', 'Neu log[0].status == "PENDING": cho them 5 giay roi goi lai (retry 1 lan).'),
    ('4', 'Neu log[0].status == "SENT": thay toast -> "Email da gui den {recipient}"  (mau xanh la).'),
    ('5', 'Neu log[0].status == "FAILED": thay toast -> "Gui email that bai"  (mau do).'),
    ('6', 'Neu khong co log (khach khong co email): khong hien toast email.'),
])
api_box(doc, [
    ('Endpoint', 'GET /api/notifications/log'),
    ('Params poll', 'referenceId={reservationId}&size=1'),
    ('Response path', 'res.data.data[0]  -- PageResponse<NotificationLogDto> truc tiep (KHONG qua ApiResponse)'),
    ('NotificationLogDto fields', 'id, type, channel, recipient, template, status (PENDING/SENT/FAILED), errorMessage, referenceId, referenceType, sentAt'),
])

h2(doc, '3.2  Chuong thong bao (Bell)')
body(doc,
     'Bieu tuong chuong (icon Bell) tren header cua man hinh WAITER. '
     'Khi nhan vao, goi GET /api/notifications/log?size=20 de lay 20 ban ghi moi nhat. '
     'Hien thi danh sach: template, recipient, status, thoi gian tuong doi.')
api_box(doc, [
    ('Endpoint', 'GET /api/notifications/log'),
    ('Params bell', 'size=20'),
    ('Response path', 'res.data.data[]  -- PageResponse<NotificationLogDto> truc tiep'),
    ('Trigger', 'bellOpen state thay doi thanh true (useEffect([bellOpen]))'),
])
note(doc, 'Quyen truy cap: WAITER / CASHIER / MANAGER / ADMIN. WAITER chi query theo referenceId cu the nen khong lo bi lo thong tin.', color=SUCCESS)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. CASHIER -- DANH SACH BAN (PHAN KET NOI BE)
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '4. Cashier -- Phan da ket noi BE')
role_line(doc, ['CASHIER', 'MANAGER'])
body(doc,
     'Route: /cashier. File: src/components/cashier/CashierOrders.tsx. '
     'Luu y: chi phan danh sach ban la du lieu that. Danh sach mon an va tao order la mock/chua lam.')

h2(doc, '4.1  Load danh sach ban')
body(doc, 'Component mount -> useEffect goi listTables().')
api_box(doc, [
    ('Endpoint', 'GET /api/tables'),
    ('Response path', 'res.data.data[]  -- ApiResponse<TableDto[]>'),
    ('Su dung', 'Hien thi so do ban tren man hinh chinh, mau sac theo status'),
])
body(doc, 'Mau ban theo trang thai:')
bul(doc, 'AVAILABLE: mau xanh la -- "Trong, san sang"')
bul(doc, 'OCCUPIED: mau do -- "Dang co khach"')
bul(doc, 'BILLING: mau cam -- "Cho thanh toan"')
bul(doc, 'RESERVED: mau vang -- "Da dat truoc"')
bul(doc, 'CLEANING: mau xam -- "Dang don dep"')
add_screenshot(doc, '09_cashier_tables.png', 'Man hinh Cashier -- so do ban (du lieu that tu API)')

h2(doc, '4.2  Phan CHUA ket noi BE (de lam sau)')
bul(doc, 'Danh sach mon an: MENU_ITEMS la mang hardcode 6 item ten "Item Name" 89.000d')
bul(doc, 'Danh muc mon: CATEGORIES hardcode (All, Ban Chay, Maki, Ura Maki, Special Rolls, Nuoc Uong)')
bul(doc, 'Tao order: nut bam co trong UI nhung CHUA goi API nao de submit order')
bul(doc, 'ORDER_ITEMS hien thi: hardcode 3 item "Dang nau"')

h2(doc, '4.3  Dang nhap / Doi mat khau / Dang xuat Cashier')
body(doc, 'Da ket noi BE (xem chuong 1): Cashier Header co menu avatar -> "Doi mat khau" va "Dang xuat".')
api_box(doc, [
    ('Doi mat khau', 'POST /api/auth/change-password (xem muc 1.3)'),
    ('Dang xuat', 'POST /api/auth/logout (xem muc 1.4)'),
])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. NHAT KY THAO TAC (AUDIT LOG)
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '5. Nhat ky Thao tac (Audit Log)')
role_line(doc, ['MANAGER', 'ADMIN'])
body(doc,
     'Route: /manager/audit-logs. File: src/components/audit/AuditLogPage.tsx. '
     'API: src/api/auditLogs.ts.')

h2(doc, '5.1  Load du lieu va bo loc')
body(doc,
     'useEffect([actorUsername, action, targetEntity, from, to]) -> useCallback load(page) -> '
     'goi listAuditLogs(). Moi thay doi bo loc tu dong trigger load(0) (reset ve trang 1).')
api_box(doc, [
    ('Endpoint', 'GET /api/audit-logs'),
    ('Params', 'actorUsername?, action?, targetEntity?, targetId?, from? (YYYY-MM-DD), to? (YYYY-MM-DD), page=0, size=30'),
    ('Response path', 'res.data.data[]  va  res.data.pagination  -- PageResponse<AuditLogResponse> truc tiep (KHONG qua ApiResponse)'),
    ('Sort', 'createdAt DESC (BE da sort co dinh)'),
    ('Quyen', 'hasAnyRole("MANAGER","ADMIN") -- WAITER/CASHIER nhan 403'),
])

add_screenshot(doc, '12_audit_log.png', 'Trang Nhat ky thao tac -- bo loc va danh sach')

h2(doc, '5.2  Cac truong hien thi')
body(doc, 'AuditLogResponse: id, actorId, actorUsername, action, targetEntity, targetId, detail (JSON string), ipAddress, createdAt (ISO).')
body(doc, 'FE hien thi:')
bul(doc, 'Thoi gian: format dd/MM/yyyy HH:mm:ss')
bul(doc, 'Nhan vien: actorUsername')
bul(doc, 'Hanh dong: badge mau theo ACTION_META (xem muc 5.3)')
bul(doc, 'Doi tuong: ENTITY_LABELS[targetEntity] (Reservation -> "Dat ban", User -> "Nhan vien")')
bul(doc, 'ID: targetId.slice(0,8) + "..." (hover xem day du)')
bul(doc, 'Chi tiet: JSON.parse(detail) -> hien key:value, nhan mo rong xem het')

h2(doc, '5.3  Mapping action -> nhan hien thi')
body(doc, '16 action duoc dinh nghia trong ACTION_META (AuditLogPage.tsx):')

rows_am = [
    ('RESERVATION_CREATE',       'Tao dat ban',           'mau xanh duong #025cca'),
    ('RESERVATION_CONFIRM',      'Xac nhan dat ban',      'mau xanh duong'),
    ('RESERVATION_ASSIGN_TABLE', 'Xep ban',               'mau xanh duong'),
    ('RESERVATION_CHECK_IN',     'Check-in',              'mau xanh la #0d9e6e'),
    ('RESERVATION_NO_SHOW',      'Khong den',             'mau xam #888'),
    ('RESERVATION_CANCEL',       'Huy dat ban',           'mau do #e53935'),
    ('RESERVATION_UPDATE',       'Sua dat ban',           'mau xanh duong'),
    ('USER_CREATE',              'Tao nhan vien',         'mau tim #7c3aed'),
    ('USER_UPDATE',              'Sua nhan vien',         'mau tim'),
    ('USER_DELETE',              'Xoa nhan vien',         'mau do'),
    ('USER_UNLOCK',              'Mo khoa TK',            'mau xanh la'),
    ('AUTH_LOGIN',               'Dang nhap',             'mau cam #e67e00'),
    ('AUTH_LOGIN_FAILED',        'Dang nhap that bai',    'mau do'),
    ('AUTH_LOGOUT',              'Dang xuat',             'mau cam'),
    ('AUTH_PASSWORD_CHANGED',    'Doi mat khau',          'mau cam'),
    ('AUTH_PASSWORD_RESET',      'Reset mat khau',        'mau cam'),
    ('AUTH_ACCOUNT_ACTIVATED',   'Kich hoat TK',          'mau xanh la'),
]

t = doc.add_table(rows=1 + len(rows_am), cols=3)
t.style = 'Table Grid'; t.autofit = False
t.columns[0].width = Cm(5.5); t.columns[1].width = Cm(4.5); t.columns[2].width = Cm(5)
hdr = t.rows[0].cells
for cell, label in zip(hdr, ['Action (key BE)', 'Nhan hien thi', 'Mau badge']):
    shd(cell, '025CCA')
    p = no_space_para(cell)
    r = p.add_run(label); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5)
for i, (action, lbl, color) in enumerate(rows_am):
    row = t.rows[i + 1].cells
    bg = 'F5F6F7' if i % 2 == 0 else 'FFFFFF'
    for cell in row: shd(cell, bg)
    p0 = no_space_para(row[0]); r0 = p0.add_run(action); r0.font.size = Pt(9); r0.font.name = 'Consolas'
    p1 = no_space_para(row[1]); r1 = p1.add_run(lbl); r1.font.size = Pt(9.5)
    p2 = no_space_para(row[2]); r2 = p2.add_run(color); r2.font.size = Pt(9.5)
doc.add_paragraph()

h2(doc, '5.4  Bo loc va phan trang')
bul(doc, 'Bo loc Nhan vien: input text, loc theo actorUsername chua chuoi (param: actorUsername=).')
bul(doc, 'Bo loc Hanh dong: dropdown 17 option (label tieng Viet map tu ACTION_META), param: action=.')
bul(doc, 'Bo loc Doi tuong: dropdown Reservation / User, param: targetEntity=.')
bul(doc, 'Tu ngay / Den ngay: input type=date, gui dinh dang YYYY-MM-DD, param: from= & to=.')
bul(doc, 'Nut "Xoa bo loc": reset tat ca state -> useCallback tu dong re-fetch.')
bul(doc, 'Phan trang: 30 ban ghi/trang. Nut <- Truoc / Tiep ->, hien "Trang X / Y" va "Tong Z ban ghi".')
add_screenshot(doc, '13_audit_log_filtered.png', 'Nhat ky thao tac -- da ap dung bo loc Hanh dong')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. DASHBOARD -- WIDGET HOAT DONG GAN DAY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '6. Dashboard -- Widget "Hoat dong gan day"')
role_line(doc, ['MANAGER', 'ADMIN'])
body(doc,
     'Route: /manager/dashboard. File: src/components/dashboard/RecentActivities.tsx. '
     'Day LA PHAN DUY NHAT tren Dashboard dung du lieu that tu BE. '
     'Phan con lai (KPI cards, bieu do, theo doi nhan vien) dung mockData.')

h2(doc, '6.1  Luong load')
step_tbl(doc, [
    ('1', 'Component mount -> useEffect goi listAuditLogs({ size: 10 })'),
    ('2', 'Goi: GET /api/audit-logs?size=10   (lay 10 hanh dong moi nhat, sort DESC createdAt)'),
    ('3', 'setLogs(r.data.data)  -- truong data cua PageResponse'),
    ('4', 'Hien thi danh sach trong Card "Hoat dong gan day".'),
])
api_box(doc, [
    ('Endpoint', 'GET /api/audit-logs'),
    ('Params', 'size=10'),
    ('Response path', 'r.data.data[]  -- PageResponse<AuditLogResponse>.data'),
    ('Luu y', 'KHONG co ApiResponse wrapper -- endpoint tra PageResponse truc tiep'),
])

h2(doc, '6.2  Hien thi moi dong')
bul(doc, 'Icon mau theo nhom: calendar (RESERVATION*), nguoi (USER*), khoa (AUTH*)')
bul(doc, 'Ten nhan vien: actorUsername (dam xanh)')
bul(doc, '"vua" + hanh dong tieng Viet tu ACTION_LABEL map')
bul(doc, 'Thoi gian tuong doi: ham timeAgo(createdAt) -- "Vua xong" / "X phut truoc" / "X gio truoc" / "X ngay truoc"')
note(doc, 'Luu y: widget chi load 1 lan khi mount, KHONG tu refresh. F5 hoac navigate lai de xem cap nhat.', color=ORANGE)
add_screenshot(doc, '11_dashboard_recent_activities.png', 'Widget "Hoat dong gan day" tren Dashboard Manager')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. PHU LUC -- TONG HOP API DA DUNG
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '7. Phu luc -- Tong hop API da duoc goi tu FE')

body(doc, 'Toan bo request deu qua src/api/apiClient.ts -- axios instance, baseURL: "/api", tu dong gan Bearer token.')
body(doc, '')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'; t.autofit = False
t.columns[0].width = Cm(1.5)
t.columns[1].width = Cm(5.5)
t.columns[2].width = Cm(4.0)
t.columns[3].width = Cm(4.0)
for cell, lbl in zip(t.rows[0].cells, ['Method', 'Endpoint', 'Goi tu file', 'Muc dich']):
    shd(cell, '025CCA')
    p = no_space_para(cell); r = p.add_run(lbl)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)

api_rows = [
    ('POST',   '/api/auth/login',                  'auth.ts / LoginPage',       'Dang nhap'),
    ('POST',   '/api/auth/verify/info',             'auth.ts / LoginPage',       'Gui OTP kich hoat TK'),
    ('POST',   '/api/auth/verify/otp',              'auth.ts / LoginPage',       'Xac thuc OTP lan dau'),
    ('POST',   '/api/auth/resend-otp',              'auth.ts / LoginPage',       'Gui lai OTP'),
    ('POST',   '/api/auth/forgot-password',         'auth.ts / ForgotPasswordPage', 'Quen mat khau'),
    ('POST',   '/api/auth/reset-password',          'auth.ts / NewPasswordPage', 'Dat lai mat khau'),
    ('POST',   '/api/auth/change-password',         'auth.ts / ChangePasswordModal', 'Doi mat khau'),
    ('POST',   '/api/auth/logout',                  'auth.ts / Reservation/Cashier', 'Dang xuat'),
    ('POST',   '/api/auth/refresh',                 'apiClient.ts (interceptor)', 'Lam moi token 401'),
    ('GET',    '/api/reservations',                 'reservations.ts / Reservation', 'Lay ds dat ban'),
    ('POST',   '/api/reservations',                 'reservations.ts / ReservationModal', 'Tao dat ban'),
    ('PUT',    '/api/reservations/{id}/confirm',    'reservations.ts / Reservation', 'Xac nhan dat ban'),
    ('DELETE', '/api/reservations/{id}',            'reservations.ts / Reservation', 'Huy dat ban'),
    ('PUT',    '/api/reservations/{id}/check-in',   'reservations.ts / Reservation', 'Check-in khach'),
    ('PUT',    '/api/reservations/{id}/no-show',    'reservations.ts / Reservation', 'Danh dau khong den'),
    ('PUT',    '/api/reservations/{id}',            'reservations.ts / Reservation', 'Xep ban (body: tableId)'),
    ('GET',    '/api/tables',                       'tables.ts / Reservation/Cashier/Modal', 'Lay ds ban'),
    ('GET',    '/api/notifications/log',            'notifications.ts / Reservation', 'Poll ket qua email'),
    ('GET',    '/api/notifications/log',            'notifications.ts / Reservation', 'Lich su email (bell)'),
    ('GET',    '/api/audit-logs',                   'auditLogs.ts / AuditLogPage', 'Trang audit log'),
    ('GET',    '/api/audit-logs',                   'auditLogs.ts / RecentActivities', 'Widget hoat dong'),
]

for i, (method, endpoint, src, purpose) in enumerate(api_rows):
    row = t.add_row().cells
    bg = 'F5F6F7' if i % 2 == 0 else 'FFFFFF'
    for cell in row: shd(cell, bg)
    method_colors = {'GET': SUCCESS, 'POST': PRIMARY, 'PUT': ORANGE, 'DELETE': DANGER}
    p0 = no_space_para(row[0]); r0 = p0.add_run(method)
    r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = method_colors.get(method, DARK)
    p1 = no_space_para(row[1]); r1 = p1.add_run(endpoint)
    r1.font.size = Pt(8.5); r1.font.name = 'Consolas'
    p2 = no_space_para(row[2]); r2 = p2.add_run(src); r2.font.size = Pt(9)
    p3 = no_space_para(row[3]); r3 = p3.add_run(purpose); r3.font.size = Pt(9)

doc.add_paragraph()

# footer
doc.add_paragraph()
fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('WRMS -- Chi tai lieu tinh nang da ket noi Backend  |  Phien ban 1.1  |  Thang 6/2026')
r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True

# save
out = 'docs/WRMS_Huong_Dan_Su_Dung.docx'
doc.save(out)
print(f'OK  Saved -> {out}')
